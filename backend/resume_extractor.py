# resume_extractor.py
import os
import tempfile
from pathlib import Path
from typing import Tuple
import requests
from pypdf import PdfReader
from docx import Document


class ResumeExtractor:
    """Extract text from various resume formats."""
    
    @staticmethod
    def extract_from_pdf(file_path: str) -> str:
        """Extract text from PDF file."""
        try:
            pdf_reader = PdfReader(file_path)
            text = ""
            
            for page in pdf_reader.pages:
                text += page.extract_text()
            
            return text.strip()
        except Exception as e:
            raise ValueError(f"Error extracting PDF: {str(e)}")
    
    @staticmethod
    def extract_from_docx(file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables if any
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
            
            return text.strip()
        except Exception as e:
            raise ValueError(f"Error extracting DOCX: {str(e)}")
    
    @staticmethod
    def extract_from_google_docs(doc_url: str) -> str:
        """
        Extract text from Google Docs by converting to PDF first.
        Google Docs share link should be publicly accessible.
        """
        try:
            # Convert Google Docs to PDF export URL
            if "docs.google.com/document" in doc_url:
                # Extract document ID from URL
                doc_id = doc_url.split("/d/")[1].split("/")[0]
                pdf_url = f"https://docs.google.com/document/d/{doc_id}/export?format=pdf"
                
                # Download PDF
                response = requests.get(pdf_url, timeout=10)
                response.raise_for_status()
                
                # Save to temporary file
                with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_file:
                    tmp_file.write(response.content)
                    tmp_path = tmp_file.name
                
                # Extract text from PDF
                text = ResumeExtractor.extract_from_pdf(tmp_path)
                
                # Clean up
                os.unlink(tmp_path)
                
                return text
            else:
                raise ValueError("Invalid Google Docs URL")
        except Exception as e:
            raise ValueError(f"Error extracting from Google Docs: {str(e)}")
    
    @staticmethod
    def extract_from_file(file_path: str) -> str:
        """
        Extract text from file based on extension.
        """
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == ".pdf":
            return ResumeExtractor.extract_from_pdf(file_path)
        elif file_extension in [".docx", ".doc"]:
            return ResumeExtractor.extract_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    @staticmethod
    def extract_from_url(url: str) -> Tuple[str, str]:
        """
        Extract text from URL (Google Docs or file URL).
        Returns tuple of (text, source_type)
        """
        if "docs.google.com" in url:
            text = ResumeExtractor.extract_from_google_docs(url)
            return text, "google_docs"
        elif url.endswith((".pdf", ".docx", ".doc")):
            try:
                response = requests.get(url, timeout=10)
                response.raise_for_status()
                
                # Save to temporary file
                with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
                    tmp_file.write(response.content)
                    tmp_path = tmp_file.name
                
                text = ResumeExtractor.extract_from_file(tmp_path)
                os.unlink(tmp_path)
                
                return text, "file_url"
            except Exception as e:
                raise ValueError(f"Error downloading file from URL: {str(e)}")
        else:
            raise ValueError("Unsupported URL format")